"""
Report Generator module for Farm Management System
Generates comprehensive reports in DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List, Dict, Any
import os

class ReportGenerator:
    """Generates various reports in DOCX format"""
    
    def __init__(self, output_dir: str = "reports"):
        """Initialize report generator"""
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    def _create_document(self, title: str) -> Document:
        """Create a new document with standard formatting"""
        doc = Document()
        
        # Add title
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph()
        date_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Spacing
        
        return doc
    
    def _save_document(self, doc: Document, filename: str) -> str:
        """Save document and return filepath"""
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath
    
    def generate_animal_profile(self, animal_data: Dict[str, Any], health_records: List[Dict], 
                               vaccinations: List[Dict], weight_records: List[Dict]) -> str:
        """Generate comprehensive animal profile report"""
        doc = self._create_document(f"Animal Profile: {animal_data.get('tag_number', 'Unknown')}")
        
        # Basic Information Section
        doc.add_heading('Basic Information', 1)
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'
        
        basic_info = [
            ('Tag Number', animal_data.get('tag_number', '')),
            ('Name', animal_data.get('name', '')),
            ('Species', animal_data.get('species', '')),
            ('Breed', animal_data.get('breed', '')),
            ('Gender', animal_data.get('gender', '')),
            ('Date of Birth', animal_data.get('date_of_birth', '')),
            ('Current Weight', f"{animal_data.get('weight', 'N/A')} kg" if animal_data.get('weight') else 'N/A'),
            ('Color', animal_data.get('color', '')),
            ('Status', animal_data.get('status', '')),
            ('Acquisition Date', animal_data.get('acquisition_date', '')),
        ]
        
        for i, (label, value) in enumerate(basic_info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        # Health Records Section
        doc.add_paragraph()
        doc.add_heading('Health Records', 1)
        
        if health_records:
            for record in health_records[:10]:  # Last 10 records
                doc.add_heading(f"Record Date: {record.get('record_date', 'N/A')}", 2)
                p = doc.add_paragraph()
                p.add_run(f"Type: ").bold = True
                p.add_run(f"{record.get('record_type', 'N/A')}\n")
                
                if record.get('diagnosis'):
                    p.add_run(f"Diagnosis: ").bold = True
                    p.add_run(f"{record.get('diagnosis')}\n")
                
                if record.get('treatment'):
                    p.add_run(f"Treatment: ").bold = True
                    p.add_run(f"{record.get('treatment')}\n")
                
                if record.get('medications'):
                    p.add_run(f"Medications: ").bold = True
                    p.add_run(f"{record.get('medications')}\n")
                
                if record.get('veterinarian'):
                    p.add_run(f"Veterinarian: ").bold = True
                    p.add_run(f"{record.get('veterinarian')}\n")
        else:
            doc.add_paragraph("No health records found.")
        
        # Vaccination Records Section
        doc.add_paragraph()
        doc.add_heading('Vaccination Records', 1)
        
        if vaccinations:
            vac_table = doc.add_table(rows=1, cols=4)
            vac_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = vac_table.rows[0].cells
            header_cells[0].text = 'Vaccine Name'
            header_cells[1].text = 'Date'
            header_cells[2].text = 'Next Due'
            header_cells[3].text = 'Veterinarian'
            
            for vac in vaccinations:
                row_cells = vac_table.add_row().cells
                row_cells[0].text = vac.get('vaccine_name', '')
                row_cells[1].text = vac.get('vaccination_date', '')
                row_cells[2].text = vac.get('next_due_date', '')
                row_cells[3].text = vac.get('veterinarian', '')
        else:
            doc.add_paragraph("No vaccination records found.")
        
        # Weight History Section
        doc.add_paragraph()
        doc.add_heading('Weight History', 1)
        
        if weight_records:
            weight_table = doc.add_table(rows=1, cols=3)
            weight_table.style = 'Light Grid Accent 1'
            
            header_cells = weight_table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Weight (kg)'
            header_cells[2].text = 'Body Condition Score'
            
            for record in weight_records[-20:]:  # Last 20 records
                row_cells = weight_table.add_row().cells
                row_cells[0].text = record.get('weight_date', '')
                row_cells[1].text = str(record.get('weight', ''))
                row_cells[2].text = str(record.get('body_condition_score', 'N/A'))
        else:
            doc.add_paragraph("No weight records found.")
        
        # Save and return
        filename = f"animal_profile_{animal_data.get('tag_number', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return self._save_document(doc, filename)
    
    def generate_herd_report(self, animals: List[Dict[str, Any]]) -> str:
        """Generate comprehensive herd report"""
        doc = self._create_document("Herd Management Report")
        
        # Summary Statistics
        doc.add_heading('Herd Summary', 1)
        
        total_animals = len(animals)
        species_count = {}
        gender_count = {}
        status_count = {}
        
        for animal in animals:
            species = animal.get('species', 'Unknown')
            species_count[species] = species_count.get(species, 0) + 1
            
            gender = animal.get('gender', 'Unknown')
            gender_count[gender] = gender_count.get(gender, 0) + 1
            
            status = animal.get('status', 'Unknown')
            status_count[status] = status_count.get(status, 0) + 1
        
        doc.add_paragraph(f"Total Animals: {total_animals}")
        doc.add_paragraph()
        
        # Species breakdown
        doc.add_heading('By Species', 2)
        for species, count in species_count.items():
            doc.add_paragraph(f"{species}: {count}", style='List Bullet')
        
        # Gender breakdown
        doc.add_heading('By Gender', 2)
        for gender, count in gender_count.items():
            doc.add_paragraph(f"{gender}: {count}", style='List Bullet')
        
        # Status breakdown
        doc.add_heading('By Status', 2)
        for status, count in status_count.items():
            doc.add_paragraph(f"{status}: {count}", style='List Bullet')
        
        # Detailed Animal List
        doc.add_page_break()
        doc.add_heading('Detailed Animal List', 1)
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Tag Number', 'Name', 'Species', 'Breed', 'Gender', 'Age/DOB', 'Status']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Data rows
        for animal in animals:
            row_cells = table.add_row().cells
            row_cells[0].text = animal.get('tag_number', '')
            row_cells[1].text = animal.get('name', '')
            row_cells[2].text = animal.get('species', '')
            row_cells[3].text = animal.get('breed', '')
            row_cells[4].text = animal.get('gender', '')
            row_cells[5].text = animal.get('date_of_birth', '')
            row_cells[6].text = animal.get('status', '')
        
        filename = f"herd_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return self._save_document(doc, filename)
    
    def generate_health_report(self, health_records: List[Dict[str, Any]], 
                              start_date: str = None, end_date: str = None) -> str:
        """Generate health management report"""
        doc = self._create_document("Health Management Report")
        
        if start_date and end_date:
            doc.add_paragraph(f"Period: {start_date} to {end_date}")
        
        doc.add_heading('Health Records Summary', 1)
        doc.add_paragraph(f"Total Records: {len(health_records)}")
        
        # Categorize by record type
        record_types = {}
        for record in health_records:
            rec_type = record.get('record_type', 'Unknown')
            record_types[rec_type] = record_types.get(rec_type, 0) + 1
        
        doc.add_heading('Records by Type', 2)
        for rec_type, count in record_types.items():
            doc.add_paragraph(f"{rec_type}: {count}", style='List Bullet')
        
        # Detailed records
        doc.add_page_break()
        doc.add_heading('Detailed Health Records', 1)
        
        for record in health_records:
            doc.add_heading(f"Date: {record.get('record_date', 'N/A')}", 2)
            p = doc.add_paragraph()
            p.add_run(f"Animal ID: ").bold = True
            p.add_run(f"{record.get('animal_id', 'N/A')}\n")
            p.add_run(f"Type: ").bold = True
            p.add_run(f"{record.get('record_type', 'N/A')}\n")
            
            if record.get('diagnosis'):
                p.add_run(f"Diagnosis: ").bold = True
                p.add_run(f"{record.get('diagnosis')}\n")
            
            if record.get('treatment'):
                p.add_run(f"Treatment: ").bold = True
                p.add_run(f"{record.get('treatment')}\n")
            
            if record.get('medications'):
                p.add_run(f"Medications: ").bold = True
                p.add_run(f"{record.get('medications')} ({record.get('dosage', '')})\n")
        
        filename = f"health_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return self._save_document(doc, filename)
    
    def generate_financial_report(self, financial_records: List[Dict[str, Any]], 
                                 start_date: str = None, end_date: str = None) -> str:
        """Generate financial report"""
        doc = self._create_document("Financial Report")
        
        if start_date and end_date:
            doc.add_paragraph(f"Period: {start_date} to {end_date}")
        
        # Calculate totals
        total_income = sum(r.get('amount', 0) for r in financial_records if r.get('transaction_type') == 'Income')
        total_expenses = sum(r.get('amount', 0) for r in financial_records if r.get('transaction_type') == 'Expense')
        net = total_income - total_expenses
        
        doc.add_heading('Financial Summary', 1)
        summary_table = doc.add_table(rows=3, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_table.rows[0].cells[0].text = 'Total Income'
        summary_table.rows[0].cells[1].text = f"${total_income:,.2f}"
        summary_table.rows[1].cells[0].text = 'Total Expenses'
        summary_table.rows[1].cells[1].text = f"${total_expenses:,.2f}"
        summary_table.rows[2].cells[0].text = 'Net Profit/Loss'
        summary_table.rows[2].cells[1].text = f"${net:,.2f}"
        
        # Category breakdown
        doc.add_paragraph()
        doc.add_heading('Expenses by Category', 2)
        
        expense_categories = {}
        for record in financial_records:
            if record.get('transaction_type') == 'Expense':
                category = record.get('category', 'Uncategorized')
                expense_categories[category] = expense_categories.get(category, 0) + record.get('amount', 0)
        
        for category, amount in expense_categories.items():
            doc.add_paragraph(f"{category}: ${amount:,.2f}", style='List Bullet')
        
        # Detailed transactions
        doc.add_page_break()
        doc.add_heading('Detailed Transactions', 1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Date'
        header_cells[1].text = 'Type'
        header_cells[2].text = 'Category'
        header_cells[3].text = 'Amount'
        header_cells[4].text = 'Description'
        
        for record in financial_records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.get('transaction_date', '')
            row_cells[1].text = record.get('transaction_type', '')
            row_cells[2].text = record.get('category', '')
            row_cells[3].text = f"${record.get('amount', 0):,.2f}"
            row_cells[4].text = record.get('description', '')
        
        filename = f"financial_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return self._save_document(doc, filename)
    
    def generate_breeding_report(self, breeding_records: List[Dict[str, Any]]) -> str:
        """Generate breeding management report"""
        doc = self._create_document("Breeding Management Report")
        
        doc.add_heading('Breeding Summary', 1)
        doc.add_paragraph(f"Total Breeding Records: {len(breeding_records)}")
        
        successful = sum(1 for r in breeding_records if r.get('success'))
        doc.add_paragraph(f"Successful Breedings: {successful}")
        doc.add_paragraph(f"Success Rate: {(successful/len(breeding_records)*100):.1f}%" if breeding_records else "N/A")
        
        # Detailed records
        doc.add_page_break()
        doc.add_heading('Detailed Breeding Records', 1)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Breeding Date', 'Dam ID', 'Sire ID', 'Method', 'Expected Delivery', 'Success']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        for record in breeding_records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.get('breeding_date', '')
            row_cells[1].text = str(record.get('dam_id', ''))
            row_cells[2].text = str(record.get('sire_id', ''))
            row_cells[3].text = record.get('breeding_method', '')
            row_cells[4].text = record.get('expected_delivery', '')
            row_cells[5].text = 'Yes' if record.get('success') else 'No'
        
        filename = f"breeding_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return self._save_document(doc, filename)
